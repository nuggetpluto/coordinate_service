from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def _set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold

    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)

    _set_run_font(
        run,
        size=14 if level == 1 else 12,
        bold=True,
        color=(31, 78, 121)
    )

    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def _add_text(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_font(run, size=11)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def _add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    _set_run_font(run, size=11)


def _format_value(value):
    if isinstance(value, float):
        return f"{value:.3f}"
    return str(value)


def _add_simple_table(document, df, columns_map):
    table = document.add_table(rows=1, cols=len(columns_map))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    for index, title in enumerate(columns_map.values()):
        header_cells[index].text = title

    for _, row in df.iterrows():
        row_cells = table.add_row().cells

        for index, column in enumerate(columns_map.keys()):
            row_cells[index].text = _format_value(row[column])

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, size=9)

    return table


def generate_report(
        source_system,
        target_system,
        input_df,
        result_df,
        output_path,
        transform_parameters=None
):
    document = Document()

    for section in document.sections:
        section.top_margin = Pt(50)
        section.bottom_margin = Pt(50)
        section.left_margin = Pt(55)
        section.right_margin = Pt(55)

    _add_heading(document, "Отчёт по преобразованию координат", level=1)

    _add_heading(document, "1. Введение", level=2)
    _add_text(
        document,
        "В этом отчёте представлены результаты преобразования координат "
        "между выбранными геодезическими системами координат."
    )

    _add_heading(document, "2. Параметры ввода", level=2)

    _add_bullet(document, f"Исходная таблица данных: {len(input_df)} точек")
    _add_bullet(document, f"Начальная система: {source_system}")
    _add_bullet(document, f"Конечная система: {target_system}")

    if transform_parameters:
        parameters_text = (
            f"Параметры: "
            f"dx: {transform_parameters['dx']}, "
            f"dy: {transform_parameters['dy']}, "
            f"dz: {transform_parameters['dz']}, "
            f"wx: {transform_parameters['wx']}, "
            f"wy: {transform_parameters['wy']}, "
            f"wz: {transform_parameters['wz']}, "
            f"m: {transform_parameters['m']}"
        )
        _add_bullet(document, parameters_text)

    _add_heading(
        document,
        "3. Общая формула перехода между выбранными системами",
        level=2
    )

    formula = (
        "(X, Y, Z)ᵦ = (1 + m) · R · (X, Y, Z)ₐ + "
        "(ΔX, ΔY, ΔZ)"
    )

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(formula)
    _set_run_font(run, size=12, bold=True)

    _add_heading(
        document,
        "4. Таблицы с координатами до и после преобразования",
        level=2
    )

    _add_text(document, "Координаты до преобразований")

    _add_simple_table(
        document,
        input_df,
        {
            "Name": "Имя",
            "X": "Начальная X",
            "Y": "Начальная Y",
            "Z": "Начальная Z"
        }
    )

    document.add_paragraph()

    _add_text(document, "Координаты после преобразований")

    _add_simple_table(
        document,
        result_df,
        {
            "Name": "Имя",
            "X_new": "Конечная X",
            "Y_new": "Конечная Y",
            "Z_new": "Конечная Z"
        }
    )

    _add_heading(document, "5. Вывод", level=2)
    _add_text(
        document,
        "Процесс преобразования координат был успешно выполнен. "
        "Полученные значения представлены в таблицах выше и могут быть "
        "использованы для дальнейшего анализа или экспорта."
    )

    document.save(output_path)